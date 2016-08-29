# -*- coding: utf-8 -*-
# Create program with options:
# Create report from directory(specify directory)
# Create forms from report(specify report file)
# Recolor matrices(specify directory)
# Update BNPP Training(specify Training file and directory to update from)
import tkFileDialog
import tkMessageBox
from Tkinter import *
import os
import docx
#tkFileDialog.askopenfile()
#tkFileDialog.askdirectory()
#root=Tkinter.Tk()
#root.withdraw()

def parsename(name):
	if "xlsx" in name:
		cut=5
	elif "xls" in name:
		cut=4
	new=""
	t=0
	j=0
	for x in range(0,len(name)):
		if name[x]=="-":
			t=1
		if t==1 and name[x]==" ":
			t=2
			print name[x+1]
		elif t==2 and name[x]==" ":
			#print name[x-1]
			j=x+1
			#print name[j]
			break
	return name[j:-cut]

def pop_up_message(label,message):
	top=Tk()
	def box(label,message):
		label=label
		message=message
		tkMessageBox.showinfo(label,message)
	B1=Button(top,text=label)
	B1.pack()

def specify_file():
	root2=Tk()
	root2.withdraw()
	return tkFileDialog.askopenfilename()
	
def specify_directory():
	root2=Tk()
	root2.withdraw()
	return tkFileDialog.askdirectory()

# def menus():
    # no_selection_is_made=1
    # choices="""
    # 1 Create report from directory(specify directory)
    # 2 Create forms from report(specify report file)
    # 3 Recolor matrices(specify directory)
    # 4 Update BNPP Training(specify Training file and directory to update from)
    # """
    # while no_selection_is_made:
        # print "Select a choice."
        # try:
            # choice=int(raw_input(choices))
        # except:
            # print "invalid choice"
        # if choice==1:
            # return create_report()
        # elif choice==2:
            # return create_forms()
        # elif choice==3:
            # return recolor_matrices()
        # elif choice==4:
            # return bnpp_update()
        # else:
            # print "invalid choice"
def create_report():
	directory=specify_directory()
	compilation=new_wkbk()
	all_wkbks_2=[]
	for sub,_,files in os.walk(directory):
		for file in files:
			if "Matrix" in file:
				all_wkbks_2.append(sub+file)
	try:
		all_wkbks_2.remove(compilation)
	except:pass
	for workbook in all_wkbks_2:
		workbook=parsename(workbook)
		active_wkbk(compilation)
		wbn=workbook[:10]
		new_sheet(wbn)
		#active_wkbk(workbook)
		labels=["Reference","Rev Number","Rev Date","Training Date","Frequency Required"]
		if Cell(7,3).is_empty():
			for column in range(0,5):
				Cell(wbn,1,(column+1)).value=labels[column]
			counter=6
			counter2=2
			active_wkbk(workbook)
			a=0
			while Cell(counter,1).value!="Review date":
				if not Cell(counter,1).is_empty():
					active_wkbk(workbook)
					g=0
					revnum=[]
					revdate=[]
					reference="'"+Cell(counter,1).value
					if "QAP" in reference:
						reference="'"+reference[4:]
					revno=Cell(counter,4).value
					if revno is not None:
						for x in revno:
							if g==0:
								revnum.append(x)
							elif g==1:
								revdate.append(x)
							if x==" ":
								g=1
					revnum=revnum[:-1]
					revnum="".join(revnum)
					revdate="".join(revdate)
					trainingdate=Cell(counter,5).value
					frequency=Cell(counter,6).value
					active_wkbk(compilation)
					active_sheet(wbn)
					Cell(wbn,counter2-a,1).value=reference
					Cell(wbn,counter2-a,2).value=revnum
					Cell(wbn,counter2-a,3).value=revdate
					Cell(wbn,counter2-a,4).value=trainingdate
					Cell(wbn,counter2-a,5).value=frequency
				else:
					a+=1
				counter2+=1
				counter+=1
				active_wkbk(workbook)
		else:
			for column in range(0,5):
				Cell(wbn,1,(column+1)).value=labels[column]
				counter=6
				counter2=2
				active_wkbk(workbook)
				a=0
			while Cell(counter,1).value!="Review date":
				if not Cell(counter,1).is_empty():
					active_wkbk(workbook)
					g=0
					revnum=[]
					revdate=[]
					reference="'"+Cell(counter,1).value
				if "QAP" in reference:
					reference="'"+reference[4:]
					revno=Cell(counter,3).value
				if revno is not None:
					for x in revno:
						if g==0:
							revnum.append(x)
						elif g==1:
							revdate.append(x)
						if x==" ":
							g=1
					revnum=revnum[:-1]
					revnum="".join(revnum)
					revdate="".join(revdate)
					trainingdate=Cell(counter,4).value
					frequency=Cell(counter,5).value
					active_wkbk(compilation)
					active_sheet(wbn)
					Cell(wbn,counter2-a,1).value=reference
					Cell(wbn,counter2-a,2).value=revnum
					Cell(wbn,counter2-a,3).value=revdate
					Cell(wbn,counter2-a,4).value=trainingdate
					Cell(wbn,counter2-a,5).value=frequency
				else:
					a+=1
				counter2+=1
				counter+=1
				active_wkbk(workbook)
	pop()
def create_report_from_f():
	file=specify_file()
	compilation=new_wkbk()
	all_wkbks_2=[]
	for sub,_,files in os.walk(directory):
		for file in files:
			if "Matrix" in file:
				all_wkbks_2.append(sub+file)
	try:
		all_wkbks_2.remove(compilation)
	except:pass
	workbook=parsename(file)
	active_wkbk(compilation)
	wbn=workbook[:10]
	new_sheet(wbn)
	#active_wkbk(workbook)
	labels=["Reference","Rev Number","Rev Date","Training Date","Frequency Required"]
	if Cell(7,3).is_empty():
		for column in range(0,5):
			Cell(wbn,1,(column+1)).value=labels[column]
		counter=6
		counter2=2
		active_wkbk(workbook)
		a=0
		while Cell(counter,1).value!="Review date":
			if not Cell(counter,1).is_empty():
				active_wkbk(workbook)
				g=0
				revnum=[]
				revdate=[]
				reference="'"+Cell(counter,1).value
				if "QAP" in reference:
					reference="'"+reference[4:]
				revno=Cell(counter,4).value
				if revno is not None:
					for x in revno:
						if g==0:
							revnum.append(x)
						elif g==1:
							revdate.append(x)
						if x==" ":
							g=1
				revnum=revnum[:-1]
				revnum="".join(revnum)
				revdate="".join(revdate)
				trainingdate=Cell(counter,5).value
				frequency=Cell(counter,6).value
				active_wkbk(compilation)
				active_sheet(wbn)
				Cell(wbn,counter2-a,1).value=reference
				Cell(wbn,counter2-a,2).value=revnum
				Cell(wbn,counter2-a,3).value=revdate
				Cell(wbn,counter2-a,4).value=trainingdate
				Cell(wbn,counter2-a,5).value=frequency
			else:
				a+=1
			counter2+=1
			counter+=1
			active_wkbk(workbook)
	else:
		for column in range(0,5):
			Cell(wbn,1,(column+1)).value=labels[column]
			counter=6
			counter2=2
			active_wkbk(workbook)
			a=0
		while Cell(counter,1).value!="Review date":
			if not Cell(counter,1).is_empty():
				active_wkbk(workbook)
				g=0
				revnum=[]
				revdate=[]
				reference="'"+Cell(counter,1).value
			if "QAP" in reference:
				reference="'"+reference[4:]
				revno=Cell(counter,3).value
			if revno is not None:
				for x in revno:
					if g==0:
						revnum.append(x)
					elif g==1:
						revdate.append(x)
					if x==" ":
						g=1
				revnum=revnum[:-1]
				revnum="".join(revnum)
				revdate="".join(revdate)
				trainingdate=Cell(counter,4).value
				frequency=Cell(counter,5).value
				active_wkbk(compilation)
				active_sheet(wbn)
				Cell(wbn,counter2-a,1).value=reference
				Cell(wbn,counter2-a,2).value=revnum
				Cell(wbn,counter2-a,3).value=revdate
				Cell(wbn,counter2-a,4).value=trainingdate
				Cell(wbn,counter2-a,5).value=frequency
			else:
				a+=1
			counter2+=1
			counter+=1
			active_wkbk(workbook)
	pop()
def QAM_forms(namedac,positiondac,directory,training_form,date):
	def dictsa():
		namedict={}
		posdict={}
		for line in namedac:
			name=[]
			value=0
			t=0
			for x in str(line):
				if x==":":
					t=1
				elif t==0:
					name.append(x)
				elif t==1:
					try:
						value=int(x)
					except:
						pass
			name=''.join(name)
			namedict[name]=value
		for line in positiondac:
			name=[]
			position=[]
			t=0
			for x in line:
				if x==":":
					t=1
				elif t==0:
					name.append(x)
				elif t==1:
					position.append(x)
			name=''.join(name)
			position=''.join(position)
			posdict[name]=position[:-1]
		return namedict,posdict
	nd,ps=dictsa()
	for name in nd:
		t=0
		doc1=docx.Document(training_form)
		try:
			doc1.styles.add_style('Procedures',docx.enum.style.WD_STYLE_TYPE.CHARACTER)
		except:
			pass
		try:
			doc1.styles.add_style('Nor',docx.enum.style.WD_STYLE_TYPE.CHARACTER)
		except:
			pass
		try:
			doc1.styles.add_style('Underlined',docx.enum.style.WD_STYLE_TYPE.CHARACTER)
		except:pass
		try:
			doc1.styles.add_style('Bolded',docx.enum.style.WD_STYLE_TYPE.CHARACTER)
		except:pass
		und=doc1.styles['Procedures']
		und.font.name='Arial'
		und.font.size=docx.shared.Pt(11)
		und.font.underline=True
		und.font.bold=True
		nor=doc1.styles['Nor']
		nor.font.name='Arial'
		nor.font.size=docx.shared.Pt(11)
		under=doc1.styles['Underlined']
		under.font.name='Arial'
		under.font.size=docx.shared.Pt(11)
		under.font.underline=True
		under=doc1.styles['Bolded']
		under.font.name='Arial'
		under.font.size=docx.shared.Pt(11)
		under.font.bold=True
		for paragraph in doc1.paragraphs:
			if "Describe the Purpose of the Training" in paragraph.text:
				stra=""
				for x in paragraph.text:
					if x==":":
						stra+=x
						break
					else:
						stra+=x
				stra=stra+" "
				strb="Required Reading to "
				strc="QAM R8"
				strd="Section 21 Addendum Riverport Shop"
				stre="Section 22 Addendum Barakah Project"
				paragraph.clear()
				paragraph.add_run(stra,"Bolded")
				if nd[name] in [1,4,6,9]:
					paragraph.add_run(strc,'Underlined')
					t=1
				if nd[name] in [3,4,8,9]:
					if nd[name]==4:
						paragraph.add_run(" and "+ strd,'Underlined')
					elif t==1:
						paragraph.add_run(", ",'Underlined')
					else:
						paragraph.add_run(strd,'Underlined')
				if nd[name] in [5,6,8,9]:
					if t==1:
						paragraph.add_run(" and ",'Underlined')
					paragraph.add_run(stre,'Underlined')
				paragraph.add_run(" Revision Changes",'Underlined')
			if "attach or list" in paragraph.text:
				stra=""
				for x in paragraph.text:
					if x==":":
						stra+=x
						break
					else:
						stra+=x
				stra=stra+" "
				r8string="Complete QA Manual (QAM) R8"
				r1string="Complete Section 21 R1 Addendum Riverport"
				r3string="Complete Section 23 R3 Addendum Barakah Project"
				revchang="revision changes attachment"
				paragraph.clear()
				paragraph.add_run(stra,'Bolded')
				if nd[name] in [1,4,6,9]:
					paragraph.add_run(r8string,'Procedures')
					t=1
				if nd[name] in [3,4,8,9]:
					if t==1:
						paragraph.add_run(", ",'Procedures')
					paragraph.add_run(r1string,'Procedures')
					t=1
				if nd[name] in [5,6,8,9]:
					if t==1:
						paragraph.add_run(", ",'Procedures')
					paragraph.add_run(r3string,'Procedures')
				paragraph.add_run(", and "+revchang,'Procedures')
				#paragraph.add_run(strc,'Nor')
		for table in doc1.tables:
			for paragraph in table.cell(1,0).paragraphs:
				paragraph.add_run(name,'Nor')
				#table.cell(1,0).text=name
			for paragraph in table.cell(1,2).paragraphs:
				paragraph.add_run(ps[name],'Nor')
				#table.cell(1,2).text=ps[name]
		doc1.save(directory+name+" QA Manual R8 Required Reading Training Form %s.docx"%date)
def proc_forms(namedac,positiondac,directory,bnpp_file,reportfile,date,training_form):	
	def dictsa():
		namedict={}
		posdict={}
		for line in namedac:
			name=[]
			value=0
			t=0
			for x in str(line):
				if x==":":
					t=1
				elif t==0:
					name.append(x)
				elif t==1:
					try:
						value=int(x)
					except:
						pass
			name=''.join(name)
			namedict[name]=value
		for line in positiondac:
			name=[]
			position=[]
			t=0
			for x in line:
				if x==":":
					t=1
				elif t==0:
					name.append(x)
				elif t==1:
					position.append(x)
			name=''.join(name)
			position=''.join(position)
			posdict[name]=position[:-1]
		return namedict,posdict

	nd,ps=dictsa()

	#create procedure and name dict
	active_wkbk(parsename(bnpp_file))
	active_sheet("Version of Procedures")

	procdict={}

	active_wkbk(report)
	l1=2
	for sheet in all_sheets():
		active_sheet(sheet)
		active_wkbk(index[0])
		active_sheet(index[1])
		while not Cell(l1,1).is_empty():
			if "QAM" not in Cell(l1,1).value:
				procdict[Cell(l1,1).value]="QAP "+Cell(l1,1).value+" R"+str(Cell(l1,2).value)+" "+Cell(l1,4).value
			l1+=1
		active_wkbk(report)

	#decide which procedures need to be trained to
	active_wkbk(report)
	persondict={}
	for sheet in all_sheets():
		active_sheet(sheet)
		proclist=[]
		l1=2
		while not Cell(l1,1).is_empty():
			print Cell(l1,1).value,Cell(l1,8).value,Cell(l1,9).value
			#if (Cell(l1,8).value=="No" and Cell(l1,1).value != "2-5") or (Cell(l1,9).value=="No" and Cell(l1,1).value != "2-5"):
			if (Cell(l1,8).value=="No") or (Cell(l1,9).value=="No"):
				proclist.append(Cell(l1,1).value)
			l1+=1
		persondict[sheet]=proclist
		print proclist,sheet
	for person in persondict:
		p=""
		for name in nd:
			if person in name:
				p=name
		t=0
		doc1=docx.Document(training_form)
		try:
			doc1.styles.add_style('Procedures',docx.enum.style.WD_STYLE_TYPE.CHARACTER)
		except:
			pass
		try:
			doc1.styles.add_style('Nor',docx.enum.style.WD_STYLE_TYPE.CHARACTER)
		except:
			pass
		try:
			doc1.styles.add_style('Underlined',docx.enum.style.WD_STYLE_TYPE.CHARACTER)
		except:pass
		try:
			doc1.styles.add_style('Bolded',docx.enum.style.WD_STYLE_TYPE.CHARACTER)
		except:pass
		und=doc1.styles['Procedures']
		und.font.name='Arial'
		und.font.size=docx.shared.Pt(11)
		und.font.underline=True
		und.font.bold=True
		nor=doc1.styles['Nor']
		nor.font.name='Arial'
		nor.font.size=docx.shared.Pt(11)
		under=doc1.styles['Underlined']
		under.font.name='Arial'
		under.font.size=docx.shared.Pt(11)
		under.font.underline=True
		under=doc1.styles['Bolded']
		under.font.name='Arial'
		under.font.size=docx.shared.Pt(11)
		under.font.bold=True
		for paragraph in doc1.paragraphs:
			# if "Describe the Purpose of the Training" in paragraph.text:
				# stra=""
				# for x in paragraph.text:
					# if x==":":
						# stra+=x
						# break
					# else:
						# stra+=x
				# stra=stra+" "
				# strb="Required Reading to "
				# strc="QAM R8"
				# strd="Section 21 Addendum Riverport Shop"
				# stre="Section 22 Addendum Barakah Project"
				# paragraph.clear()
				# paragraph.add_run(stra,"Bolded")
				# for 
				# paragraph.add_run(" Revision Changes",'Underlined')
			if "attach or list" in paragraph.text:
				stra=""
				for x in paragraph.text:
					if x==":":
						stra+=x
						break
					else:
						stra+=x
				stra=stra+" "
				# r8string="Complete QA Manual (QAM) R8"
				# r1string="Complete Section 21 R1 Addendum Riverport"
				# r3string="Complete Section 23 R3 Addendum Barakah Project"
				# revchang="revision changes attachment"
				paragraph.clear()
				paragraph.add_run(stra,'Bolded')
				pdi=persondict[person]
				for i in range(0,len(persondict[person])):
					for proc in procdict:
						if proc==pdi[i] and i<(len(pdi)-2):
							paragraph.add_run(procdict[proc]+", ",'Procedures')
						elif proc==pdi[i] and i<(len(pdi)-1):
							paragraph.add_run(procdict[proc]+", and ",'Procedures')
						if proc==pdi[i] and i==(len(pdi)-1):
							paragraph.add_run(procdict[proc]+".",'Procedures')
				#paragraph.add_run(strc,'Nor')
		for table in doc1.tables:
			for paragraph in table.cell(1,0).paragraphs:
				paragraph.add_run(p,'Nor')
				#table.cell(1,0).text=name
			for paragraph in table.cell(1,2).paragraphs:
				paragraph.add_run(ps[p],'Nor')
				#table.cell(1,2).text=ps[name]
		doc1.save(directory+p+' Required Readings Training Form %s.docx'%date)
def pathic(path2):
	i=0
	for x in path2:
		if x=="/":
			i=0
		i+=1
	i-=1
	return path2[-i:]
def create_forms():
	print "Select names dictionary"
	names=specify_file()
	print "Select positions dictionary"
	positions=specify_file()
	ch=raw_input("For QAM(1) or for procedures(0)?")
	print "Select output directory"
	directory=specify_directory()
	print "Select form template"
	training_form=specify_file()
	date=raw_input("Input date in format MMDDYY>")
	if ch==1:
		QAM_forms(names,positions,directory,date,training_form)
	elif ch==0:
		print "Select procedure index file"
		bnpp_file=specify_file()
		print "Select report file"
		reportfile=specify_file()
		proc_forms(names,position,directory,bnpp_file,reportfile,date,training_form)
	pop()
def recolor_matrices():
	directory=specify_directory()
	#import and define help
	#create list of relevant workbooks and their paths
	def dir(directory):
		list=[]
		for sub,dir,files in os.walk(directory):
			for file in files:
				if "Training Matrix" in file and "Historic" not in sub:
					list.append([sub+"\\"+file,file])
		return list
	#color matrices
	def clsrrs():
		rows=0
		l0=1
		while not "Review" in str(Cell(l0,1).value):
			#Cell(l0,15).value=rows
			if "Reference" in str(Cell(l0,1).value):
				startrow=l0
				collength=1
				rows=0
				if Cell(startrow,3).is_empty():
					c=0
					while not Cell(startrow,collength).is_empty() or c==1:
						if c==1 and Cell(startrow,collength+1).is_empty():
							break
						elif Cell(startrow,collength+1).is_empty():
							c=1
							collength+=1
							continue
						collength+=1
				else:
					while not Cell(startrow,collength).is_empty():
						collength+=1
					collength=collength-1
			l0+=1
			rows+=1
		rows=rows-2
		return collength,rows,startrow
	#color top row
	def colortop(startrow,collength):
		for column in range(1,collength+1):
			Cell(startrow,column).color="b7dee8"
	#bold top row
	def boldtop(collength,startrow,wkbk):
		if collength==7 and startrow==5:
			pass
		elif collength==8 and startrow==5:
			pass
		else:
			txt=open("missedbolds.txt","a")
			txt.write("for workbook %s the top was not bolded"%wkbk)
			txt.close()
	#set up rows that need to be colored
	def createrowlist(startrow,rows):
		start=0
		addgo=0
		rowlist=[]
		for row in range(startrow,startrow+rows+1):
			if row>=(startrow+2) and Cell(row,1).is_empty() and start==0:
				if not Cell(row+1,1).is_empty():
					start=1
					addgo=1
			elif row>=(startrow+2) and not Cell(row,1).is_empty() and start==0:
				start=1
				addgo=1
				rowlist.append(row)
				if not Cell(row+1,1).is_empty():
					addgo=0
			elif start==1 and addgo==1:
				rowlist.append(row)
				if not Cell(row+1,1).is_empty():
					addgo=0
			elif addgo==0:
				if not Cell(row+1,1).is_empty():
					addgo=1
		return rowlist
	#Cell(1,13).value=rowlist
	def colorrows(rowlist,startrow,rows,collength):
		for row in range(startrow,startrow+rows+1):
			for column in range(1,collength+1):
				if not Cell(row,column).color=="yellow" and row in rowlist:
					Cell(row,column).color="fde9d9"
				elif not Cell(row,column).color=="yellow" and startrow!=row:
					Cell(row,column).color="white"
	#Cell(1,14).value=collength
	def main(sub,file):
			open_wkbk(sub)
			active_wkbk(file)
			collength,rows,startrow=clsrrs()
			rowlist=createrowlist(startrow,rows)
			colortop(startrow,collength)
			boldtop(collength,startrow,file)
			colorrows(rowlist,startrow,rows,collength)
			save()
			close_wkbk(file)
	#MAIN
	directory=dir()
	for set in directory:
		sub=set[0]
		file=set[1]
		main(sub,file)
		pass
			
	pop()
def recolor_matrices_from_f():
	directory=specify_directory()
	#import and define help
	#create list of relevant workbooks and their paths
	def dir(directory):
		list=[]
		for sub,dir,files in os.walk(directory):
			for file in files:
				if "Training Matrix" in file and "Historic" not in sub:
					list.append([sub+"\\"+file,file])
		return list
	#color matrices
	def clsrrs():
		rows=0
		l0=1
		while not "Review" in str(Cell(l0,1).value):
			#Cell(l0,15).value=rows
			if "Reference" in str(Cell(l0,1).value):
				startrow=l0
				collength=1
				rows=0
				if Cell(startrow,3).is_empty():
					c=0
					while not Cell(startrow,collength).is_empty() or c==1:
						if c==1 and Cell(startrow,collength+1).is_empty():
							break
						elif Cell(startrow,collength+1).is_empty():
							c=1
							collength+=1
							continue
						collength+=1
				else:
					while not Cell(startrow,collength).is_empty():
						collength+=1
					collength=collength-1
			l0+=1
			rows+=1
		rows=rows-2
		return collength,rows,startrow
	#color top row
	def colortop(startrow,collength):
		for column in range(1,collength+1):
			Cell(startrow,column).color="b7dee8"
	#bold top row
	def boldtop(collength,startrow,wkbk):
		if collength==7 and startrow==5:
			pass
		elif collength==8 and startrow==5:
			pass
		else:
			txt=open("missedbolds.txt","a")
			txt.write("for workbook %s the top was not bolded"%wkbk)
			txt.close()
	#set up rows that need to be colored
	def createrowlist(startrow,rows):
		start=0
		addgo=0
		rowlist=[]
		for row in range(startrow,startrow+rows+1):
			if row>=(startrow+2) and Cell(row,1).is_empty() and start==0:
				if not Cell(row+1,1).is_empty():
					start=1
					addgo=1
			elif row>=(startrow+2) and not Cell(row,1).is_empty() and start==0:
				start=1
				addgo=1
				rowlist.append(row)
				if not Cell(row+1,1).is_empty():
					addgo=0
			elif start==1 and addgo==1:
				rowlist.append(row)
				if not Cell(row+1,1).is_empty():
					addgo=0
			elif addgo==0:
				if not Cell(row+1,1).is_empty():
					addgo=1
		return rowlist
	#Cell(1,13).value=rowlist
	def colorrows(rowlist,startrow,rows,collength):
		for row in range(startrow,startrow+rows+1):
			for column in range(1,collength+1):
				if not Cell(row,column).color=="yellow" and row in rowlist:
					Cell(row,column).color="fde9d9"
				elif not Cell(row,column).color=="yellow" and startrow!=row:
					Cell(row,column).color="white"
	#Cell(1,14).value=collength
	def main(sub,file):
			open_wkbk(sub)
			active_wkbk(file)
			collength,rows,startrow=clsrrs()
			rowlist=createrowlist(startrow,rows)
			colortop(startrow,collength)
			boldtop(collength,startrow,file)
			colorrows(rowlist,startrow,rows,collength)
			save()
			close_wkbk(file)
	#MAIN
	file=specify_file()
	main(file,parsename(file))
	pop()
def bnpp_update():
	bnpp=specify_file()
	directory=specify_directory()
	open_wkbk(bnpp)
	active_wkbk(parsename(bnpp))
	active_sheet("Version of Procedures")

	qaplist=[]
	qapsilist=[]

	for dirpath,_,filenames in os.walk(directory):
		
		locale=0
		for x in range(0,len(dirpath)):
			if dirpath[x]=="\\":
				locale=x
				
		dr=dirpath[locale:]
		
		if dr=="\QAP":
			print "YES YES"
			for f in filenames:
				qaplist.append(f)
		elif dr=="\QAP-SI":
			for f in filenames:
				qapsilist.append(f)
	#print qaplist,qapsilist     
	l1=2
	while not Cell(l1,1).is_empty():
		t=0
		
		procedure=Cell(l1,1).value
		if "'" in procedure:
			procedure=procedure[1:]
			
		#if procedure in qaplist:
		for qap in qaplist:
			if procedure in qap:
				t=1
				Cell(l1,4).value=parsename(qap)

		#elif procedure in qapsilist:
		if t==0:
			for qapsi in qapsilist:
				if procedure in qapsi:
					Cell(l1,4).value=parsename(qapsi)

		l1+=1
	pop()
def pop():
	pop_up_message('Success','Ok')
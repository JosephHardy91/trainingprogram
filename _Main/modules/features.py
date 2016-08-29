# !/usr/bin/python
# -*- coding: utf-8 -*-
import tkFileDialog
import tkMessageBox
from Tkinter import *
import os
import docx
import datetime
import sys
# --------internal imports------------
from decisionparse import *
from utilities import *
import shutil
from cell import *
import datetime
from win32com.client import Dispatch


def lower(thing):
    if isinstance(thing, basestring):
        return thing.lower()
    else:
        return thing


# features

# ----reports----#
def create_report():
    directory = specify_directory()
    #directory = r"http://iscommunity.isco-pipe.com/QC/nuclear/Indoctrination Training Records/BNPP Project"
    compilation = new_wkbk()
    all_wkbks_2 = []
    wkbk_dict = {}
    for sub, dircs, files in os.walk(directory):
        # print "pass"
        # print sub
        for file in files:
            if "Training Matrix" in file and "Historical" not in sub:
                new_s = ""
                for x in range(0, len(sub)):
                    if sub[x] == "\\":
                        new_s += "/"
                    else:
                        new_s += sub[x]
                if parsename(file) in all_wkbks_2:
                    break
                else:
                    if ".xls" in file:
                        # print new_s+file
                        all_wkbks_2.append(new_s + "/" + file)
                        wkbk_dict[new_s + "/" + file] = file
    try:
        all_wkbks_2.remove(compilation)
    except:
        pass
    for workbook in all_wkbks_2:
        # print workbook
        open_wkbk(workbook)
        # print 2
        wbn = wkbk_dict[workbook]
        # print 3
        workbook2 = pathic(workbook)
        # print 4
        active_wkbk(compilation)
        # print 5
        wbnd = str(wbn)
        wbn = wbn[:10]
        if wbn[-1] == " ":
            wbn = wbn[:-1]
        new_sheet(wbn)
        # print wbnd, workbook

        active_wkbk(wbnd)
        print wbnd
        labels = ["Reference", "Rev Number", "Rev Date", "Training Date", "Frequency Required"]
        if Cell(7, 3).is_empty():
            active_wkbk(compilation)
            active_sheet(wbn)
            # print active_sheet()
            for column in range(0, 5):
                Cell(wbn, 1, (column + 1)).value = labels[column]
            counter = 6
            counter2 = 2
            # print wbnd
            active_wkbk(wbnd)
            # print active_sheet()
            a = 0
            # print Cell(14,1).value
            while lower(Cell(counter, 1).value) != "review date":
                # print 1,wkbkpaths[active_wkbk()], counter, Cell(counter, 1).value
                if not Cell(counter, 1).is_empty():
                    print Cell(counter, 1).value
                    active_wkbk(wbnd)
                    g = 0
                    revnum = []
                    revdate = []
                    try:
                        reference = "'" + Cell(counter, 1).value
                    except TypeError as e:
                        print ';'.join(["\n", str(counter), str(e)])
                    if "QAP" in reference:
                        if "SI" not in reference:
                            reference = "'" + reference[4:]
                        else:
                            reference = "'" + reference[5:]
                    if reference[len(reference) - 1] == "-":
                        reference = reference[:len(reference) - 1]
                    revno = Cell(counter, 4).value
                    if revno is not None:
                        for x in revno:
                            if g == 0:
                                revnum.append(x)
                            elif g == 1:
                                revdate.append(x)
                            if x == " ":
                                g = 1
                    revnum = revnum[:-1]
                    revnum = "".join(revnum)
                    revdate = "".join(revdate)
                    try:
                        revdate = (datetime.datetime.strptime(revdate, "%m/%d/%Y")).date()
                    except:
                        pass
                    trainingdate = Cell(counter, 5).value
                    frequency = Cell(counter, 6).value
                    active_wkbk(compilation)
                    active_sheet(wbn)
                    Cell(wbn, counter2 - a, 1).value = reference
                    Cell(wbn, counter2 - a, 2).value = revnum
                    Cell(wbn, counter2 - a, 3).value = revdate
                    Cell(wbn, counter2 - a, 4).value = trainingdate
                    Cell(wbn, counter2 - a, 5).value = frequency
                    # print Cell(counter2-a,4).value, Cell(counter2-a,5).value
                else:
                    a += 1
                counter2 += 1
                counter += 1
                active_wkbk(wbnd)
        else:
            active_wkbk(compilation)
            active_sheet(wbn)
            for column in range(0, 5):
                Cell(wbn, 1, (column + 1)).value = labels[column]
            counter = 6
            counter2 = 2
            # print wbnd
            active_wkbk(wbnd)
            # print active_sheet()
            a = 0
            # print Cell(14,1).value
            while Cell(counter, 1).value != "Review date":
                # print 2,wkbkpaths[active_wkbk()], counter, Cell(counter, 1).value
                if not Cell(counter, 1).is_empty():
                    active_wkbk(wbnd)
                    g = 0
                    revnum = []
                    revdate = []
                    reference = "'" + Cell(counter, 1).value
                    if "QAP" in reference:
                        if "SI" not in reference:
                            reference = "'" + reference[4:]
                        else:
                            reference = "'" + reference[5:]
                    if reference[len(reference) - 1] == "-":
                        reference = reference[:len(reference) - 1]
                    revno = Cell(counter, 3).value
                    if revno is not None:
                        for x in revno:
                            if g == 0:
                                revnum.append(x)
                            elif g == 1:
                                revdate.append(x)
                            if x == " ":
                                g = 1
                    revnum = revnum[:-1]
                    revnum = "".join(revnum)
                    revdate = "".join(revdate)
                    try:
                        revdate = (datetime.datetime.strptime(revdate, "%m/%d/%Y")).date()
                    except:
                        pass
                    trainingdate = Cell(counter, 4).value
                    frequency = Cell(counter, 5).value
                    active_wkbk(compilation)
                    active_sheet(wbn)
                    Cell(wbn, counter2 - a, 1).value = reference
                    Cell(wbn, counter2 - a, 2).value = revnum
                    Cell(wbn, counter2 - a, 3).value = revdate
                    Cell(wbn, counter2 - a, 4).value = trainingdate
                    Cell(wbn, counter2 - a, 5).value = frequency
                else:
                    a += 1
                counter2 += 1
                counter += 1
                active_wkbk(wbnd)
        close_wkbk(wbnd)
    active_wkbk(compilation)
    # print active_wkbk()
    save_path = os.path.normpath(
        os.path.join(os.path.abspath(''), '..', 'Reports', os.path.split(directory)[1] + ".xlsx"))
    save(save_path)
    allsheets = all_sheets()
    addformulas(allsheets, save_path)
    xl = Dispatch("Excel.Application")
    xl.visible = True
    wb = xl.Workbooks.Open(save_path)


def create_report_from_f():
    file = specify_file()
    compilation = new_wkbk()
    all_wkbks_2 = []
    workbook = parsename(file)
    active_wkbk(compilation)
    wbn = workbook[:10]
    new_sheet(wbn)
    # active_wkbk(workbook)
    labels = ["Reference", "Rev Number", "Rev Date", "Training Date", "Frequency Required"]
    if Cell(7, 3).is_empty():
        for column in range(0, 5):
            Cell(wbn, 1, (column + 1)).value = labels[column]
        counter = 6
        counter2 = 2
        active_wkbk(workbook)
        a = 0
        while Cell(counter, 1).value != "Review date":
            if not Cell(counter, 1).is_empty():
                active_wkbk(workbook)
                g = 0
                revnum = []
                revdate = []
                reference = "'" + Cell(counter, 1).value
                if "QAP" in reference:
                    reference = "'" + reference[4:]
                revno = Cell(counter, 4).value
                if revno is not None:
                    for x in revno:
                        if g == 0:
                            revnum.append(x)
                        elif g == 1:
                            revdate.append(x)
                        if x == " ":
                            g = 1
                revnum = revnum[:-1]
                revnum = "".join(revnum)
                revdate = "".join(revdate)
                trainingdate = Cell(counter, 5).value
                frequency = Cell(counter, 6).value
                active_wkbk(compilation)
                active_sheet(wbn)
                Cell(wbn, counter2 - a, 1).value = reference
                Cell(wbn, counter2 - a, 2).value = revnum
                Cell(wbn, counter2 - a, 3).value = revdate
                Cell(wbn, counter2 - a, 4).value = trainingdate
                Cell(wbn, counter2 - a, 5).value = frequency
            else:
                a += 1
            counter2 += 1
            counter += 1
            active_wkbk(workbook)
    else:
        for column in range(0, 5):
            Cell(wbn, 1, (column + 1)).value = labels[column]
            counter = 6
            counter2 = 2
            active_wkbk(workbook)
            a = 0
        while Cell(counter, 1).value != "Review date":
            if not Cell(counter, 1).is_empty():
                active_wkbk(workbook)
                g = 0
                revnum = []
                revdate = []
                reference = "'" + Cell(counter, 1).value
            if "QAP" in reference:
                reference = "'" + reference[4:]
                revno = Cell(counter, 3).value
            if revno is not None:
                for x in revno:
                    if g == 0:
                        revnum.append(x)
                    elif g == 1:
                        revdate.append(x)
                    if x == " ":
                        g = 1
                revnum = revnum[:-1]
                revnum = "".join(revnum)
                revdate = "".join(revdate)
                trainingdate = Cell(counter, 4).value
                frequency = Cell(counter, 5).value
                active_wkbk(compilation)
                active_sheet(wbn)
                Cell(wbn, counter2 - a, 1).value = reference
                Cell(wbn, counter2 - a, 2).value = revnum
                Cell(wbn, counter2 - a, 3).value = revdate
                Cell(wbn, counter2 - a, 4).value = trainingdate
                Cell(wbn, counter2 - a, 5).value = frequency
            else:
                a += 1
            counter2 += 1
            counter += 1
            active_wkbk(workbook)
    pop()


def backwards_compatible():
    print "Specify file to be compatible with"
    b_file = specify_file('ReportBackups')
    print "Specify report to make compatible"
    reportfile, reportwkbk = file_parse('reports')
    open_wkbk(reportfile)
    active_wkbk(reportwkbk)
    names = []
    relevant = []
    relevantnames = []
    for sheet in all_sheets():
        names.append(sheet)
    # for sub,dirc,files in os.walk(directory):
    # for file in files:
    # if sub+"/"+file not in relevant:
    # open_wkbk(sub+"/"+file)
    # active_wkbk(pathic(sub+"/"+file))
    # for sheet in all_sheets():
    # if sheet in names:
    # relevant.append(sub+"/"+file)
    # for sheet in all_sheets():
    # relevantnames.append(sheet)
    # print sub+"/"+file
    # break
    # close_wkbk(file)
    open_wkbk(b_file)
    active_wkbk(pathic(b_file))
    for sheet in all_sheets():
        if sheet in names:
            for sheet in all_sheets():
                relevantnames.append(sheet)
            print "Backup from " + b_file
            break
    # for wkbk in relevant:
    relevantdict = {}
    relevantlist = []
    active_wkbk(pathic(b_file))
    for sheet in all_sheets():
        active_sheet(sheet)
        overrelevantlist = []
        if sheet in relevantnames:
            # print sheet
            l1 = 1
            while not Cell(l1, 1).is_empty():
                # if Cell(l1,9).value is None:
                # print "Warning: Possible value misread. Recommend rerun or check"
                relevantlist = [Cell(l1, 1).value, Cell(l1, 2).value, Cell(l1, 6).value, Cell(l1, 7).value,
                                Cell(l1, 9).value, Cell(l1, 10).value, Cell(l1, 3).value]
                overrelevantlist.append(relevantlist)
                l1 += 1
            # print overrelevantlist
            relevantdict[sheet] = overrelevantlist
    close_wkbk(b_file)

    active_wkbk(reportwkbk)
    print active_wkbk_var
    for sheet in all_sheets():
        active_sheet(sheet)
        if sheet in relevantdict:
            print sheet
            # print relevantdict[sheet]
            l1 = 1
            while not Cell(l1, 1).is_empty():
                for list in relevantdict[sheet]:
                    if Cell(l1, 1).value in list and Cell(l1, 2).value in list and Cell(l1, 3).value in list:
                        if Cell(l1, 6).value == list[2] and Cell(l1, 7).value == list[3]:
                            print sheet, " ", list
                            if list[4] is not None and list[4] != "None" and list[4] != u'None':
                                Cell(l1, 9).value = str(list[4])
                            if list[5] is not None and list[5] != "None" and list[5] != u'None':
                                Cell(l1, 10).value = str(list[5])
                l1 += 1
    close_wkbk(reportfile)
    print "done"


# -------forms-------#



# def QAM_forms(namedac, positiondac, directory, training_form, date):
def QAM_forms(namedac, positiondac, directory, bnpp_file, reportfile, date, training_form):
    print namedac, positiondac, directory, bnpp_file, reportfile, date, training_form
    report = pathic(reportfile)
    summaries_dir = find_self_directory('summaries')

    # results_txt=open('%s Summary.txt'%(summaries_dir+"\\"+parsename(report)),'w')

    def dictsa(namedac, positiondac):
        namedac = open(namedac, "r")
        positiondac = open(positiondac, "r")
        namedict = {}
        posdict = {}
        for line in namedac:
            name = []
            value = 0
            t = 0
            for x in str(line):
                if x == ":":
                    t = 1
                elif t == 0:
                    name.append(x)
                elif t == 1:
                    try:
                        value = int(x)
                    except:
                        pass
            name = ''.join(name)
            namedict[name] = value
        for line in positiondac:
            name = []
            position = []
            t = 0
            for x in line:
                if x == ":":
                    t = 1
                elif t == 0:
                    name.append(x)
                elif t == 1:
                    position.append(x)
            name = ''.join(name)
            position = ''.join(position)
            posdict[name] = position[:-1]
        return namedict, posdict

    nd, ps = dictsa(namedac, positiondac)
    # print nd
    # print ps
    open_wkbk(bnpp_file)
    open_wkbk(reportfile)
    # create procedure and name dict
    active_wkbk(pathic(bnpp_file))
    active_sheet("Version of Procedures")

    procdict = {}

    active_wkbk(report, data_only=True)
    l1 = 2
    for sheet in all_sheets():
        active_sheet(sheet)
        active_wkbk(pathic(bnpp_file))
        active_sheet("Version of Procedures")
        while not Cell(l1, 1).is_empty():
            if "QAM" in Cell(l1, 1).value:
                procdict[Cell(l1, 1).value] = Cell(l1, 1).value + " R" + str(Cell(l1, 2).value) + " " + Cell(
                    l1, 4).value
            l1 += 1
        active_wkbk(report)

        active_wkbk(report, data_only=True)
    persondict = {}
    persondec = {}
    # proctodec={}
    for sheet in all_sheets():
        active_sheet(sheet)
        proclist = []
        declist = {}
        l1 = 2
        while not Cell(l1, 1).is_empty():
            print Cell(l1, 1).value, Cell(l1, 8).value, Cell(l1, 9).value
            if "ibit" in Cell(l1, 1).value:
                exhibit = str(Cell(l1, 1).value)
                exhibit = exhibit.split(" ")
                exhibit[0] = "Exhibit"
                exhibit = " ".join(exhibit)
                Cell(l1, 1).value = "'" + exhibit
            if Cell(l1, 9).is_empty():
                Cell(l1, 9).value = "No"
            # print Cell(l1,1).value,Cell(l1,8).value,Cell(l1,9).value
            if ((Cell(l1, 8).value == "No" and Cell(l1, 1).value != "2-5") or (
                            Cell(l1, 9).value == "No" and Cell(l1, 1).value != "2-5")) and not (
                                Cell(l1, 8).value == "No" and Cell(l1, 9).value == "Yes" and Cell(l1,
                                                                                                  1).value == u'2-1'):
                # if (Cell(l1,8).value=="No") or (Cell(l1,9).value=="No"):
                proclist.append(str(Cell(l1, 1).value))
                try:
                    wa = str(Cell(l1, 4).value).split(" ")[1]
                except:
                    wa = str(Cell(l1, 4).value)
                declist[Cell(l1, 1).value] = [wa, str(Cell(l1, 5).value), str(Cell(l1, 8).value),
                                              str(Cell(l1, 9).value), str(Cell(l1, 10).value)]
            l1 += 1
        persondict[sheet] = proclist
        persondec[sheet] = declist

    for person in persondict:
        p = ""
        for name in nd:
            if person in name:
                p = name
        if p == "":
            print person, "not found"
        else:
            print p
        t = 0
        doc1 = docx.Document(training_form)
        try:
            doc1.styles.add_style('Procedures', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        except:
            pass
        try:
            doc1.styles.add_style('Nor', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        except:
            pass
        try:
            doc1.styles.add_style('Underlined', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        except:
            pass
        try:
            doc1.styles.add_style('Bolded', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        except:
            pass
        und = doc1.styles['Procedures']
        und.font.name = 'Arial'
        und.font.size = docx.shared.Pt(11)
        und.font.underline = True
        und.font.bold = True
        nor = doc1.styles['Nor']
        nor.font.name = 'Arial'
        nor.font.size = docx.shared.Pt(11)
        under = doc1.styles['Underlined']
        under.font.name = 'Arial'
        under.font.size = docx.shared.Pt(11)
        under.font.underline = True
        under = doc1.styles['Bolded']
        under.font.name = 'Arial'
        under.font.size = docx.shared.Pt(11)
        under.font.bold = True
        for paragraph in doc1.paragraphs:
            if paragraph.text == "Date:":
                n = 2
                datex = [date[i:i + n] for i in range(0, len(date), n)]
                daten = datex[0] + "/" + datex[1] + "/" + "20" + datex[2]
                paragraph.clear()
                paragraph.add_run("Date: ", 'Bolded')
                paragraph.add_run(daten, 'Underlined')
                # if "Describe the Purpose of the Training" in paragraph.text:
                # stra = ""
                # for x in paragraph.text:
                # if x == ":":
                # stra += x
                # break
                # else:
                # stra += x
                # stra = stra + " "
                # strb = "Required Reading to "
                # strc = "QAM R8"
                # strd = "Section 21 Addendum Riverport Shop"
                # stre = "Section 22 Addendum Barakah Project"
                # paragraph.clear()
                # paragraph.add_run(stra, "Bolded")
                # if nd[name] in [1, 4, 6, 9]:
                # paragraph.add_run(strc, 'Underlined')
                # t = 1
                # if nd[name] in [3, 4, 8, 9]:
                # if nd[name] == 4:
                # paragraph.add_run(" and " + strd, 'Underlined')
                # elif t == 1:
                # paragraph.add_run(", ", 'Underlined')
                # else:
                # paragraph.add_run(strd, 'Underlined')
                # if nd[name] in [5, 6, 8, 9]:
                # if t == 1:
                # paragraph.add_run(" and ", 'Underlined')
                # paragraph.add_run(stre, 'Underlined')
                # paragraph.add_run(" Revision Changes", 'Underlined')
                # if "attach or list" in paragraph.text:
                # stra = ""
                # for x in paragraph.text:
                # if x == ":":
                # stra += x
                # break
                # else:
                # stra += x
                # stra = stra + " "
                # r8string = "Complete QA Manual (QAM) R8"
                # r1string = "Complete Section 21 R1 Addendum Riverport"
                # r3string = "Complete Section 23 R3 Addendum Barakah Project"
                # revchang = "revision changes attachment"
                # paragraph.clear()
                # paragraph.add_run(stra, 'Bolded')
                # if nd[name] in [1, 4, 6, 9]:
                # paragraph.add_run(r8string, 'Procedures')
                # t = 1
                # if nd[name] in [3, 4, 8, 9]:
                # if t == 1:
                # paragraph.add_run(", ", 'Procedures')
                # paragraph.add_run(r1string, 'Procedures')
                # t = 1
                # if nd[name] in [5, 6, 8, 9]:
                # if t == 1:
                # paragraph.add_run(", ", 'Procedures')
                # paragraph.add_run(r3string, 'Procedures')
                # paragraph.add_run(", and " + revchang, 'Procedures')
                # paragraph.add_run(strc,'Nor')
        for table in doc1.tables:
            for paragraph in table.cell(1, 0).paragraphs:
                paragraph.add_run(p, 'Nor')
            # table.cell(1,0).text=name
            for paragraph in table.cell(1, 2).paragraphs:
                for paragraph in table.cell(1, 2).paragraphs:
                    try:
                        paragraph.add_run(ps[p], 'Nor')
                    except:
                        pass
                        # table.cell(1,2).text=ps[name]
        spath = directory + "/" + pathic(reportfile)[:-5] + "/"
        if not os.access(spath, os.F_OK):
            os.mkdir(spath)
        doc1.save(
            spath + p + ' QA Manual R9 Required Readings Training Form %s.docx' % date)
        # doc1.save(directory + p + " QA Manual R9 Required Reading Training Form %s.docx" % date)


def proc_forms(namedac, positiondac, directory, bnpp_file, reportfile, date, training_form, yeslist):
    report = pathic(reportfile)
    summaries_dir = find_self_directory('summaries')
    results_txt = open('%s Summary.txt' % (summaries_dir + "\\" + parsename(report)), 'w')
    if len(yeslist) == 1 and yeslist[0].strip() == '':
        yeslist = []
    print yeslist

    def dictsa(namedac, positiondac):
        namedac = open(namedac, "r")
        positiondac = open(positiondac, "r")
        namedict = {}
        posdict = {}
        for line in namedac:
            name = []
            value = 0
            t = 0
            for x in str(line):
                if x == ":":
                    t = 1
                elif t == 0:
                    name.append(x)
                elif t == 1:
                    try:
                        value = int(x)
                    except:
                        pass
            name = ''.join(name)
            namedict[name] = value
        for line in positiondac:
            name = []
            position = []
            t = 0
            for x in line:
                if x == ":":
                    t = 1
                elif t == 0:
                    name.append(x)
                elif t == 1:
                    position.append(x)
            name = ''.join(name)
            position = ''.join(position)
            posdict[name] = position[:-1]
        return namedict, posdict

    nd, ps = dictsa(namedac, positiondac)
    # print nd
    # print ps
    open_wkbk(bnpp_file)
    open_wkbk(reportfile)
    # create procedure and name dict
    active_wkbk(pathic(bnpp_file))
    active_sheet("Version of Procedures")

    procdict = {}

    active_wkbk(report, data_only=True)
    l1 = 2
    for sheet in all_sheets():
        active_sheet(sheet)
        active_wkbk(pathic(bnpp_file))
        active_sheet("Version of Procedures")
        while not Cell(l1, 1).is_empty():
            if str(Cell(l1, 1).value) in yeslist or len(yeslist) == 0:
                if "QAM" not in Cell(l1, 1).value:
                    procdict[Cell(l1, 1).value] = "QAP " + Cell(l1, 1).value + " R" + str(
                        Cell(l1, 2).value) + " " + Cell(
                        l1, 4).value
                elif "QAM" in Cell(l1, 1).value:
                    procdict[Cell(l1, 1).value] = Cell(l1, 1).value + " R" + str(Cell(l1, 2).value) + " " + Cell(
                        l1, 4).value
            l1 += 1
        active_wkbk(report)

    # decide which procedures need to be trained to
    active_wkbk(report, data_only=True)
    persondict = {}
    persondec = {}
    # proctodec={}
    for sheet in all_sheets():
        active_sheet(sheet)
        proclist = []
        declist = {}
        l1 = 2
        print sheet
        while not Cell(l1, 1).is_empty():
            cur_procedure=str(Cell(l1,1).value).strip()
            print cur_procedure, Cell(l1, 8).value, Cell(l1, 9).value, cur_procedure in yeslist
            if "ibit" in Cell(l1, 1).value:
                exhibit = str(Cell(l1, 1).value)
                exhibit = exhibit.split(" ")
                exhibit[0] = "Exhibit"
                exhibit = " ".join(exhibit)
                Cell(l1, 1).value = "'" + exhibit
            if Cell(l1, 9).is_empty():
                Cell(l1, 9).value = "No"
            if ((Cell(l1, 8).value == "No" and cur_procedure != "2-5") or (
                            Cell(l1, 9).value == "No" and cur_procedure != "2-5")) and not (
                                Cell(l1, 8).value == "No" and Cell(l1, 9).value == "Yes" and str(Cell(l1,
                                                                                                      1).value) == '2-1') and (str(cur_procedure) in yeslist or len(yeslist) == 0) and cur_procedure!="15-2":# and ("QAM" not in str(
                #cur_procedure) or cur_procedure in yeslist)
                # if (Cell(l1,8).value=="No") or (Cell(l1,9).value=="No"):
                print cur_procedure
                proclist.append(str(cur_procedure))
                try:
                    wa = str(Cell(l1, 4).value).split(" ")[1]
                except:
                    wa = str(Cell(l1, 4).value)
                declist[cur_procedure] = [wa, str(Cell(l1, 5).value), str(Cell(l1, 8).value),
                                              str(Cell(l1, 9).value), str(Cell(l1, 10).value)]
            l1 += 1
        persondict[sheet] = proclist
        persondec[sheet] = declist
    # print persondec
    # print proclist,sheet
    notwritten = open("notwritten.txt", "w")
    for person in persondict:
        t_1 = 0
        p = ""
        for name in nd:
            if person in name:
                p = name
        print p, persondict[person]
        t = 0
        if len(persondict[person]) > 15:
            doc1 = docx.Document(training_form)
        else:
            doc1 = docx.Document(
                r"C:\Users\User\SyncedFolder\Nuclear Training\Training Program Management(Tree Backup)\tdpstack\forms\Exhibit QAM 2.4 Process and Procedure Training Form - Copy.docx".replace(
                    "\\", "/"))
        try:
            doc1.styles.add_style('Procedures', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        except:
            pass
        try:
            doc1.styles.add_style('Nor', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        except:
            pass
        try:
            doc1.styles.add_style('Underlined', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        except:
            pass
        try:
            doc1.styles.add_style('Bolded', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        except:
            pass
        und = doc1.styles['Procedures']
        und.font.name = 'Arial'
        und.font.size = docx.shared.Pt(11)
        und.font.underline = True
        und.font.bold = True
        nor = doc1.styles['Nor']
        nor.font.name = 'Arial'
        nor.font.size = docx.shared.Pt(11)
        under = doc1.styles['Underlined']
        under.font.name = 'Arial'
        under.font.size = docx.shared.Pt(11)
        under.font.underline = True
        under = doc1.styles['Bolded']
        under.font.name = 'Arial'
        under.font.size = docx.shared.Pt(11)
        under.font.bold = True
        for paragraph in doc1.paragraphs:
            # if "Describe the Purpose of the Training" in paragraph.text:
            # stra=""
            # for x in paragraph.text:
            # if x==":":
            # stra+=x
            # break
            # else:
            # stra+=x
            # stra=stra+" "
            # strb="Required Reading to "
            # strc="QAM R8"
            # strd="Section 21 Addendum Riverport Shop"
            # stre="Section 22 Addendum Barakah Project"
            # paragraph.clear()
            # paragraph.add_run(stra,"Bolded")
            # for
            # paragraph.add_run(" Revision Changes",'Underlined')
            if paragraph.text == "Date:":
                n = 2
                datex = [date[i:i + n] for i in range(0, len(date), n)]
                daten = datex[0] + "/" + datex[1] + "/" + "20" + datex[2]
                paragraph.clear()
                paragraph.add_run("Date: ", 'Bolded')
                paragraph.add_run(daten, 'Underlined')
            if "attach or list" in paragraph.text:
                stra = ""
                for x in paragraph.text:
                    if x == ":":
                        stra += x
                        break
                    else:
                        stra += x
                stra = stra + " "
                # r8string="Complete QA Manual (QAM) R8"
                # r1string="Complete Section 21 R1 Addendum Riverport"
                # r3string="Complete Section 23 R3 Addendum Barakah Project"
                # revchang="revision changes attachment"
                paragraph.clear()
                paragraph.add_run(stra, 'Bolded')
                paragraph.alignment = 0
                pdi = persondict[person]
                lz = 0
                directory_name = directory + "/" + pathic(reportfile)[:-5] + "/" + p + "/"
                for i in range(0, len(persondict[person])):
                    for proc in procdict:
                        # print sheet
                        if proc == pdi[i]:
                            if not os.path.exists(directory_name):
                                os.makedirs(directory_name)
                            # print directory_name
                            # if "-SI" in procdict[proc]:

                            try:
                                if "-SI" in procdict[proc]:
                                    if not os.path.exists(directory_name):
                                        os.mkdir(directory_name)
                                    # notwritten.write('C:/Users/User/SyncedFolder/Nuclear Quality Assurance Program/QAP-SI/' +
                                    #     procdict[proc] + ".pdf",
                                    #     directory_name + proc + ".pdf\n")
                                    shutil.copy(
                                        'C:/Users/User/SyncedFolder/Nuclear Quality Assurance Program/QAP-SI/' +
                                        procdict[proc] + ".pdf",
                                        directory_name + proc + ".pdf")
                                else:
                                    if not os.path.exists(directory_name):
                                        os.mkdir(directory_name)
                                    # notwritten.write('C:/Users/User/SyncedFolder/Nuclear Quality Assurance Program/QAP/' +
                                    #     procdict[proc] + ".pdf",
                                    #     directory_name + proc + ".pdf\n")
                                    shutil.copy(
                                        'C:/Users/User/SyncedFolder/Nuclear Quality Assurance Program/QAP/' +
                                        procdict[proc] + ".pdf",
                                        directory_name + proc + ".pdf")
                            except:
                                try:
                                    if "-SI" in procdict[proc]:
                                        if not os.path.exists(directory_name):
                                            os.mkdir(directory_name)
                                        # notwritten.write('C:/Users/User/SyncedFolder/Nuclear Quality Assurance Program/QAP-SI/' +
                                        #     procdict[proc] + " (FCR).pdf\n")
                                        shutil.copy(
                                            'C:/Users/User/SyncedFolder/Nuclear Quality Assurance Program/QAP-SI/' +
                                            procdict[proc] + " (FCR).pdf",
                                            directory_name + proc + ".pdf")
                                    else:
                                        if not os.path.exists(directory_name):
                                            os.mkdir(directory_name)
                                        # notwritten.write('C:/Users/User/SyncedFolder/Nuclear Quality Assurance Program/QAP/' +
                                        #     procdict[proc] + " (FCR).pdf")
                                        shutil.copy(
                                            'C:/Users/User/SyncedFolder/Nuclear Quality Assurance Program/QAP/' +
                                            procdict[proc] + " (FCR).pdf",
                                            directory_name + proc + ".pdf")
                                except:
                                    notwritten.write("%s:%s not moved from %s to %s\n" % (p, procdict[proc],
                                                                                          'C:/Users/User/SyncedFolder/Nuclear Quality Assurance Program/QAP-SI/' +
                                                                                          procdict[proc] + ".pdf",
                                                                                          directory_name + proc + ".pdf"))
                            lz += 1
                            curdeclist = persondec[person][proc]
                            # print curdeclist
                            try:
                                curdeclist[4].lower()
                            except:
                                pass
                            decisionparse(p, proc, curdeclist[0], curdeclist[2].lower(), curdeclist[3].lower(),
                                          curdeclist[1].lower(), curdeclist[4])
                        if proc == pdi[i] and i < (len(pdi) - 2):
                            paragraph.add_run("\n" + procdict[proc] + ", ", 'Procedures')
                        elif proc == pdi[i] and i < (len(pdi) - 1):
                            paragraph.add_run("\n" + procdict[proc] + ", and ", 'Procedures')
                        if proc == pdi[i] and i == (len(pdi) - 1):
                            paragraph.add_run("\n" + procdict[proc] + ".", 'Procedures')
                            # paragraph.add_run(strc,'Nor')
                        if proc == pdi[i]:
                            t_1 = 1
                            results_txt.write("%s:%s\n" % (p, procdict[proc]))
        for table in doc1.tables:
            for paragraph in table.cell(1, 0).paragraphs:
                paragraph.add_run(p, 'Nor')
            # table.cell(1,0).text=name
            for paragraph in table.cell(1, 2).paragraphs:
                try:
                    paragraph.add_run(ps[p], 'Nor')
                except:
                    pass
                    # table.cell(1,2).text=ps[name]
        if lz != 0:
            spath = directory + "/" + pathic(reportfile)[:-5] + "/"
            if not os.access(spath, os.F_OK):
                os.mkdir(spath)
            if not os.path.exists(directory_name):
                os.makedirs(directory_name)
            doc1.save(
                os.path.join(directory_name, '%s Required Readings Training Form %s.docx' % (p,date)))
        if t_1 == 1:
            results_txt.write("\n")
        # if p is not None and p is not "":
        #     try:
        #         #shutil.make_archive(directory + "/" + pathic(reportfile)[:-5] + "/" + p + "/", 'zip',
        #         #                    directory + "/" + pathic(reportfile)[:-5] + "/" + p + "/")
        #         #shutil.rmtree(directory + "/" + pathic(reportfile)[:-5] + "/" + p + "/")
        #     except:
        #         pass
    masterlist = [clearedlist, backuplist, warninglist, traintonewlist, trainagainlist, updatematrixlist]
    # writefile(masterlist, parsename(report))

    notwritten.close()
    results_txt.close()
    print "Bonus Summaries written"


def create_forms():
    where = 0
    vdir = find_self_directory('tdpstack')
    # vdir = '/'.join(
    #     (((os.getcwd()).replace("\\", "/")).split("/"))[:len(((os.getcwd()).replace("\\", "/")).split("/")) - 1])
    for sub, dir, file in os.walk(vdir):
        if "tdpstack" in sub:
            sub = sub.split("\\")
            for div in sub:
                if "tdpstack" in div:
                    break
                where += 1
            # print where
            tdpdirectory = "/".join(sub[:where + 1])
            break
    if where > 0:
        stackch = "yes"
        names, positions, directory, reportfile, training_form, bnpp_file = tdpstack(tdpdirectory, "forms")
        ch = raw_input("For QAM(1) or for procedures(0)?")
        date = proper_date()
    # date=raw_input("Input date in format MMDDYY>")
    else:
        stackch = raw_input("Is there a tdp stack? Yes or No>")
        if stackch.lower() == "yes":
            names, positions, directory, reportfile, training_form, bnpp_file = tdpstack(specify_directory(), "forms")
            ch = raw_input("For QAM(1) or for procedures(0)?")
            date = raw_input("Input date in format MMDDYY>")
        elif stackch.lower() == "no":
            print "Select names dictionary"
            names = specify_file()
            print "Select positions dictionary"
            positions = specify_file()
            ch = raw_input("For QAM(1) or for procedures(0)?")
            print "Select output directory"
            directory = specify_directory()
            print "Select form template"
            training_form = specify_file()
            date = raw_input("Input date in format MMDDYY>")
    try:
        ch = int(ch)
    except:
        print "incorrect choice"
    if ch == 1:
        QAM_forms(names, positions, directory, bnpp_file, reportfile, date, training_form)
        # QAM_forms(names, positions, directory, training_form, date)
    elif ch == 0:
        if stackch.lower() == "no":
            print "Select procedure index file"
            bnpp_file = str(specify_file())
            print "Select report file"
            reportfile = str(specify_file())
        yeslist = raw_input(
            "Enter procedures in the format '3-2' or '4-3-SI',separated by commas. If you want all, leave blank").split(
            ",")
        proc_forms(names, positions, directory, bnpp_file, reportfile, date, training_form, yeslist)
    del clearedlist[:]
    del backuplist[:]
    del warninglist[:]
    del traintonewlist[:]
    del trainagainlist[:]
    del updatematrixlist[:]
    print "done"


# -----recolor_matrices----#
def recolor_matrices():
    directory = specify_directory()

    # import and define help
    # create list of relevant workbooks and their paths
    def dir(directory):
        list = []
        for sub, dir, files in os.walk(directory):
            for file in files:
                if "Training Matrix" in file and "Historic" not in sub:
                    list.append([sub + "/" + file, file])
        return list

    # color matrices
    def clsrrs():
        rows = 0
        l0 = 1
        while not "Review" in str(Cell(l0, 1).value):
            # Cell(l0,15).value=rows
            if "Reference" in str(Cell(l0, 1).value):
                startrow = l0
                collength = 1
                rows = 0
                if Cell(startrow, 3).is_empty():
                    c = 0
                    while not Cell(startrow, collength).is_empty() or c == 1:
                        if c == 1 and Cell(startrow, collength + 1).is_empty():
                            break
                        elif Cell(startrow, collength + 1).is_empty():
                            c = 1
                            collength += 1
                            continue
                        collength += 1
                else:
                    while not Cell(startrow, collength).is_empty():
                        collength += 1
                    collength = collength - 1
            l0 += 1
            rows += 1
        rows = rows - 2
        return collength, rows, startrow

    # color top row
    def colortop(startrow, collength):
        for column in range(1, collength + 1):
            Cell(startrow, column).color = "b7dee8"

    # bold top row
    def boldtop(collength, startrow, wkbk):
        if collength == 7 and startrow == 5:
            pass
        elif collength == 8 and startrow == 5:
            pass
        else:
            txt = open("missedbolds.txt", "a")
            txt.write("for workbook %s the top was not bolded" % wkbk)
            txt.close()

    # set up rows that need to be colored
    def createrowlist(startrow, rows):
        start = 0
        addgo = 0
        rowlist = []
        for row in range(startrow, startrow + rows + 1):
            if row >= (startrow + 2) and Cell(row, 1).is_empty() and start == 0:
                if not Cell(row + 1, 1).is_empty():
                    start = 1
                    addgo = 1
            elif row >= (startrow + 2) and not Cell(row, 1).is_empty() and start == 0:
                start = 1
                addgo = 1
                rowlist.append(row)
                if not Cell(row + 1, 1).is_empty():
                    addgo = 0
            elif start == 1 and addgo == 1:
                rowlist.append(row)
                if not Cell(row + 1, 1).is_empty():
                    addgo = 0
            elif addgo == 0:
                if not Cell(row + 1, 1).is_empty():
                    addgo = 1
        return rowlist

    # Cell(1,13).value=rowlist
    def colorrows(rowlist, startrow, rows, collength):
        for row in range(startrow, startrow + rows + 1):
            for column in range(1, collength + 1):
                if not Cell(row, column).color == "yellow" and row in rowlist:
                    Cell(row, column).color = "fde9d9"
                elif not Cell(row, column).color == "yellow" and startrow != row:
                    Cell(row, column).color = "white"

    # Cell(1,14).value=collength
    def main(sub, file):
        open_wkbk(sub)
        active_wkbk(file)
        collength, rows, startrow = clsrrs()
        rowlist = createrowlist(startrow, rows)
        colortop(startrow, collength)
        boldtop(collength, startrow, file)
        colorrows(rowlist, startrow, rows, collength)
        save()
        close_wkbk(file)

    # MAIN
    directory = dir()
    for set in directory:
        sub = set[0]
        file = set[1]
        main(sub, file)
        pass

    pop()


def recolor_matrices_from_f():
    file = specify_file()

    # import and define help
    # create list of relevant workbooks and their paths
    # color matrices
    def clsrrs():
        rows = 0
        l0 = 1
        while not "Review" in str(Cell(l0, 1).value):
            # Cell(l0,15).value=rows
            if "Reference" in str(Cell(l0, 1).value):
                startrow = l0
                collength = 1
                rows = 0
                if Cell(startrow, 3).is_empty():
                    c = 0
                    while not Cell(startrow, collength).is_empty() or c == 1:
                        if c == 1 and Cell(startrow, collength + 1).is_empty():
                            break
                        elif Cell(startrow, collength + 1).is_empty():
                            c = 1
                            collength += 1
                            continue
                        collength += 1
                else:
                    while not Cell(startrow, collength).is_empty():
                        collength += 1
                    collength = collength - 1
            l0 += 1
            rows += 1
        rows = rows - 2
        return collength, rows, startrow

    # color top row
    def colortop(startrow, collength):
        for column in range(1, collength + 1):
            Cell(startrow, column).color = "b7dee8"

    # bold top row
    def boldtop(collength, startrow, wkbk):
        if collength == 7 and startrow == 5:
            pass
        elif collength == 8 and startrow == 5:
            pass
        else:
            txt = open("missedbolds.txt", "a")
            txt.write("for workbook %s the top was not bolded" % wkbk)
            txt.close()

    # set up rows that need to be colored
    def createrowlist(startrow, rows):
        start = 0
        addgo = 0
        rowlist = []
        for row in range(startrow, startrow + rows + 1):
            if row >= (startrow + 2) and Cell(row, 1).is_empty() and start == 0:
                if not Cell(row + 1, 1).is_empty():
                    start = 1
                    addgo = 1
            elif row >= (startrow + 2) and not Cell(row, 1).is_empty() and start == 0:
                start = 1
                addgo = 1
                rowlist.append(row)
                if not Cell(row + 1, 1).is_empty():
                    addgo = 0
            elif start == 1 and addgo == 1:
                rowlist.append(row)
                if not Cell(row + 1, 1).is_empty():
                    addgo = 0
            elif addgo == 0:
                if not Cell(row + 1, 1).is_empty():
                    addgo = 1
        return rowlist

    # Cell(1,13).value=rowlist
    def colorrows(rowlist, startrow, rows, collength):
        for row in range(startrow, startrow + rows + 1):
            for column in range(1, collength + 1):
                if not Cell(row, column).color == "yellow" and row in rowlist:
                    Cell(row, column).color = "fde9d9"
                elif not Cell(row, column).color == "yellow" and startrow != row:
                    Cell(row, column).color = "white"

    # Cell(1,14).value=collength
    def main(sub, file):
        open_wkbk(sub)
        active_wkbk(file)
        collength, rows, startrow = clsrrs()
        rowlist = createrowlist(startrow, rows)
        colortop(startrow, collength)
        boldtop(collength, startrow, file)
        colorrows(rowlist, startrow, rows, collength)
        save()
        close_wkbk(file)

    # MAIN
    file = specify_file()
    main(file, pathic(file))
    pop()


# -----update bnpp training------#
def bnpp_update():
    print "Specify location of procedure index(BNPPTraining)"
    bnpp = specify_file('general')
    print "Specify location of directory with QAP/QAP-SI/etc folders(Nuclear Quality Assurance in Anchor)"
    directory = specify_directory()
    open_wkbk(bnpp)
    active_wkbk(pathic(bnpp))
    active_sheet("Version of Procedures")

    qaplist = []
    qapsilist = []

    for dirpath, _, filenames in os.walk(directory):
        new_dirpath = ""
        for x in dirpath:
            if x == "\\":
                new_dirpath += "/"
            else:
                new_dirpath += x
        dirpath = new_dirpath
        # print dirpath
        locale = 0
        for x in range(0, len(dirpath)):
            if dirpath[x] == "/":
                locale = x

        dr = dirpath[locale:]
        # print dr
        if dr == "/QAP":
            for f in filenames:
                if "QAP" in f:
                    qaplist.append(f)
        elif dr == "/QAP-SI":
            for f in filenames:
                if "QAP" in f:
                    qapsilist.append(f)
    # print qaplist,qapsilist
    l1 = 2
    while not Cell(l1, 1).is_empty():
        t = 0

        procedure = Cell(l1, 1).value
        if "'" in procedure:
            procedure = procedure[1:]

        # if procedure in qaplist:
        for qap in qaplist:
            if procedure in qap:
                t = 1
                number, rev, name = parseprocname(qap)
                if Cell(l1, 2).value != rev and Cell(l1, 1).value == number:
                    # print Cell(l1, 1).value, Cell(l1, 2).value, rev, qap
                    Cell(l1, 4).value = name
                    Cell(l1, 3).value = ""
                    Cell(l1, 2).value = rev

        # elif procedure in qapsilist:
        if t == 0:
            for qapsi in qapsilist:
                # print qapsi
                if procedure in qapsi:
                    number, rev, name = parseprocname(qapsi)
                    if Cell(l1, 2).value != rev and Cell(l1, 1).value == number:
                        print Cell(l1, 1).value, Cell(l1, 2).value, rev
                        Cell(l1, 4).value = name
                        Cell(l1, 3).value = ""
                        Cell(l1, 2).value = rev

        l1 += 1
    save()
    print "Successful"


# -----create_indices----
def create_file_index():
    directory = specify_directory()
    compilation = new_wkbk()
    all_wkbks_2 = []
    wkbk_dict = {}
    for sub, dircs, files in os.walk(directory):
        # print "pass"
        # print sub
        for file in files:
            if "Training Matrix" in file and "Historical" not in sub:
                new_s = ""
                for x in range(0, len(sub)):
                    if sub[x] == "\\":
                        new_s += "/"
                    else:
                        new_s += sub[x]
                if parsename(file) in all_wkbks_2:
                    break
                else:
                    if ".xls" in file:
                        # print new_s+file
                        all_wkbks_2.append(new_s + "/" + file)
                        wkbk_dict[new_s + "/" + file] = file
    try:
        all_wkbks_2.remove(compilation)
    except:
        pass
    for workbook in all_wkbks_2:
        # print workbook
        open_wkbk(workbook)
        # print 2
        wbn = wkbk_dict[workbook]
        # print 3
        workbook2 = pathic(workbook)
        # print 4
        active_wkbk(compilation)
        # print 5
        wbnd = str(wbn)
        wbn = ' '.join(wbnd.split(" ")[:2])
        new_sheet(wbn)
        # print wbnd, workbook

        active_wkbk(wbnd)
        print wbnd
        labels = ["Reference", "Rev Number", "File", "Index"]
        if Cell(7, 3).is_empty():
            active_wkbk(compilation)
            active_sheet(wbn)
            # print active_sheet()
            Cell(wbn, 1, 1).value = "Person:"
            Cell(wbn, 1, 2).value = ' '.join(wbnd.split(" ")[:2])
            for column in range(0, 4):
                Cell(wbn, 2, (column + 1)).value = labels[column]
            counter = 6
            counter2 = 3
            # print wbnd
            active_wkbk(wbnd)
            # print active_sheet()
            a = 0
            # print Cell(14,1).value
            while Cell(counter, 1).value != "Review date":
                print Cell(counter, 1).value, counter
                # print 1,wkbkpaths[active_wkbk()], counter, Cell(counter, 1).value
                # print Cell(counter, 1).value
                active_wkbk(wbnd)
                g = 0
                revnum = []
                revdate = []
                if not Cell(counter, 1).is_empty():
                    reference = "'" + Cell(counter, 1).value
                if "QAP" in reference:
                    if "SI" not in reference:
                        reference = "'" + reference[4:]
                    else:
                        reference = "'" + reference[5:]
                if reference[len(reference) - 1] == "-":
                    reference = reference[:len(reference) - 1]
                if isinstance(reference, unicode) or isinstance(reference, basestring):
                    # print "instanced"
                    while reference[0] == " " or reference[0] == "'" or reference[len(reference) - 1] == " " or (
                                ("QAM 22" in reference or "QAM 21" in reference) and (
                                            reference != "QAM 22" and reference != "QAM 21")):
                        # print "in loop"
                        if reference[0] == " " or reference[0] == "'":
                            reference = reference[1:len(reference)]
                        if reference[len(reference) - 1] == " ":
                            reference = reference[:len(reference) - 1]
                        if "QAM 22" in reference and reference != u'%s' % "QAM 22":
                            reference = u'%s' % "QAM 22"
                        elif "QAM 21" in reference and reference != u'%s' % "QAM 21":
                            reference = u'%s' % "QAM 21"
                    if "SI" in reference:
                        for x in range(0, len(reference)):
                            if x + 1 != len(reference) and reference[x] + reference[x + 1] == "SI" and reference[
                                        x - 1] != "-":
                                reference = reference[:x - 1] + "-" + reference[x:]
                revno = Cell(counter, 4).value
                if revno is not None:
                    for x in revno:
                        if g == 0:
                            revnum.append(x)
                        elif g == 1:
                            revdate.append(x)
                        if x == " ":
                            g = 1
                revnum = revnum[:-1]
                revnum = "".join(revnum)
                revdate = "".join(revdate)
                try:
                    revdate = (datetime.datetime.strptime(revdate, "%m/%d/%Y")).date()
                except:
                    pass
                trainingdate = Cell(counter, 5).value
                frequency = Cell(counter, 6).value
                active_wkbk(compilation)
                active_sheet(wbn)
                Cell(wbn, counter2 - a, 1).value = reference
                Cell(wbn, counter2 - a, 2).value = revnum
                Cell(wbn, counter2 - a,
                     3).value = '=HYPERLINK(VLOOKUP(D3,Index!A:B,2,FALSE),CONCATENATE(A3," ",B3," Training Form"))'
                # print Cell(counter2-a,4).value, Cell(counter2-a,5).value
                counter += 1
                counter2 += 1
                active_wkbk(wbnd)
                # else:
                #     a += 1
                # counter2 += 1
                # counter += 1

        else:
            active_wkbk(compilation)
            active_sheet(wbn)
            Cell(wbn, 1, 1).value = "Person:"
            Cell(wbn, 1, 2).value = matrixtoname(wbnd)
            for column in range(0, 4):
                Cell(wbn, 1, (column + 1)).value = labels[column]
            counter = 6
            counter2 = 3
            # print wbnd
            active_wkbk(wbnd)
            # print active_sheet()
            a = 0
            # print Cell(14,1).value
            while Cell(counter, 1).value != "Review date":
                print Cell(counter, 1).value, counter
                # print 2,wkbkpaths[active_wkbk()], counter, Cell(counter, 1).value
                active_wkbk(wbnd)
                g = 0
                revnum = []
                revdate = []
                if not Cell(counter, 1).is_empty():
                    reference = "'" + Cell(counter, 1).value
                if "QAP" in reference:
                    if "SI" not in reference:
                        reference = "'" + reference[4:]
                    else:
                        reference = "'" + reference[5:]
                if reference[len(reference) - 1] == "-":
                    reference = reference[:len(reference) - 1]
                    # print reference
                if isinstance(reference, unicode) or isinstance(reference, basestring):
                    # print "instanced"
                    while reference[0] == " " or reference[0] == "'" or reference[len(reference) - 1] == " " or (
                                ("QAM 22" in reference or "QAM 21" in reference) and (
                                            reference != "QAM 22" and reference != "QAM 21")):
                        # print "in loop"
                        if reference[0] == " " or reference[0] == "'":
                            reference = reference[1:len(reference)]
                        if reference[len(reference) - 1] == " ":
                            reference = reference[:len(reference) - 1]
                        if "QAM 22" in reference and reference != u'%s' % "QAM 22":
                            reference = u'%s' % "QAM 22"
                        elif "QAM 21" in reference and reference != u'%s' % "QAM 21":
                            reference = u'%s' % "QAM 21"
                    if "SI" in reference:
                        for x in range(0, len(reference)):
                            if x + 1 != len(reference) and reference[x] + reference[x + 1] == "SI" and reference[
                                        x - 1] != "-":
                                reference = reference[:x - 1] + "-" + reference[x:]
                revno = Cell(counter, 3).value
                if revno is not None:
                    for x in revno:
                        if g == 0:
                            revnum.append(x)
                        elif g == 1:
                            revdate.append(x)
                        if x == " ":
                            g = 1
                revnum = revnum[:-1]
                revnum = "".join(revnum)
                revdate = "".join(revdate)
                try:
                    revdate = (datetime.datetime.strptime(revdate, "%m/%d/%Y")).date()
                except:
                    pass
                trainingdate = Cell(counter, 4).value
                frequency = Cell(counter, 5).value
                active_wkbk(compilation)
                active_sheet(wbn)
                Cell(wbn, counter2 - a, 1).value = reference
                Cell(wbn, counter2 - a, 2).value = revnum
                Cell(wbn, counter2 - a,
                     3).value = '=HYPERLINK(VLOOKUP(D3,Index!A:B,2,FALSE),CONCATENATE(A3," ",B3," Training Form"))'
                counter2 += 1
                counter += 1
                active_wkbk(wbnd)
                # else:
                #     a += 1

        close_wkbk(wbnd)
    active_wkbk(compilation)
    # print active_wkbk()
    save_path = save_path = os.path.normpath(
        os.path.join(os.path.abspath(''), '..', 'Indices', os.path.split(directory)[1] + ".xlsx"))
    save(save_path)
